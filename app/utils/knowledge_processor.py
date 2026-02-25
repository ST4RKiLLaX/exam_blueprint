import os
import gzip
import json
import time
import ipaddress
import socket
import logging
from typing import Iterable, Iterator, List
from urllib.parse import urlparse

import faiss
import numpy as np
import pypdf
from bs4 import BeautifulSoup
from docx import Document
import pickle
import requests
from openai import OpenAI
from dotenv import load_dotenv

try:
    import tiktoken  # type: ignore
except ImportError:
    tiktoken = None

EMBEDDING_MODEL = "text-embedding-3-large"
DEFAULT_EMBEDDING_DIM = 3072
CHUNKS_JSON_GZ = "chunks.json.gz"
LEGACY_CHUNKS_GZ = "chunks.pkl.gz"
LEGACY_CHUNKS_RAW = "chunks.pkl"
MAX_REDIRECTS = 5

logger = logging.getLogger(__name__)

load_dotenv()

def create_embedding(text: str, provider: str = "openai", model: str = None) -> np.ndarray:
    """
    Create embedding using specified provider.
    
    Args:
        text: Text to embed
        provider: Provider name ("openai" or "gemini")
        model: Optional specific model name (uses provider default if None)
        
    Returns:
        numpy array of embedding vector
    """
    from app.config.provider_config import get_provider_metadata
    provider_metadata = get_provider_metadata(provider)
    
    if provider == "openai":
        from app.config.api_config import get_openai_api_key
        client = OpenAI(api_key=get_openai_api_key())
        model = model or provider_metadata.get("default_embedding_model", EMBEDDING_MODEL)
        response = client.embeddings.create(input=text, model=model)
        return np.array(response.data[0].embedding, dtype="float32")
    
    elif provider == "gemini":
        from app.utils.gemini_client import GeminiClient
        client = GeminiClient()
        model = model or provider_metadata.get("default_embedding_model", "text-embedding-004")
        response = client.embed_content(model=model, content=text)
        return np.array(response.data[0]["embedding"], dtype="float32")
    
    raise ValueError(f"Unsupported embedding provider: {provider}")

# Initialize OpenAI client with API key from config
def _get_openai_client():
    """Get OpenAI client with proper API key"""
    try:
        from app.config.api_config import get_openai_api_key
        api_key = get_openai_api_key()
        if not api_key:
            raise ValueError("No API key configured")
        return OpenAI(api_key=api_key)
    except ImportError:
        raise ValueError("API config not available")

# Don't create client at import time - create when needed

def extract_text_from_docx(file_path):
    """Stream paragraphs from DOCX file"""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue
            yield paragraph_text
        if not doc.paragraphs:
            print(f"[WARN] No paragraphs found in DOCX: {file_path}")
    except Exception as e:
        print(f"Error extracting text from DOCX (paragraph stream) at {file_path}: {e}")
        return

def extract_text_from_pdf(file_path):
    """Stream page text from PDF file to avoid large in-memory buffers"""
    try:
        with open(file_path, "rb") as file:
            pdf_reader = pypdf.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            for page_number, page in enumerate(pdf_reader.pages, start=1):
                try:
                    page_text = page.extract_text() or ""
                except Exception as page_error:
                    print(f"Error extracting text from PDF {file_path} page {page_number}: {page_error}")
                    continue
                cleaned_text = page_text.strip()
                if cleaned_text:
                    yield cleaned_text
                if page_number % 25 == 0 or page_number == total_pages:
                    print(f"📄 Processed {page_number}/{total_pages} PDF pages from {file_path}")
    except Exception as e:
        print(f"Error streaming text from PDF {file_path}: {e}")
        return

def fetch_content_from_url(url):
    """Fetch content from URL (HTML or JSON)"""
    is_valid, error_message = validate_url_for_ssrf(url)
    if not is_valid:
        raise ValueError(error_message)

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        session = requests.Session()
        session.max_redirects = MAX_REDIRECTS
        response = session.get(url, headers=headers, timeout=30, allow_redirects=True)
        if len(response.history) > MAX_REDIRECTS:
            raise ValueError("Too many redirects")
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '').lower()
        
        if 'application/json' in content_type:
            # Handle JSON content
            json_data = response.json()
            return json.dumps(json_data, indent=2)
        elif 'text/html' in content_type:
            # Handle HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            # Get text content
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            return text
        else:
            # Try to get text content anyway
            return response.text
    except Exception as e:
        print(f"Error fetching content from URL: {e}")
        return ""


def _host_is_allowlisted(hostname: str) -> bool:
    raw = os.environ.get("ALLOWED_INTERNAL_DOMAINS", "")
    allowlist = [item.strip().lower() for item in raw.split(",") if item.strip()]
    if not allowlist:
        return False
    hostname = (hostname or "").lower()
    return hostname in allowlist


def _is_blocked_ip(ip_text: str) -> bool:
    ip = ipaddress.ip_address(ip_text)
    return (
        ip.is_private
        or ip.is_loopback
        or ip.is_link_local
        or ip.is_multicast
        or ip.is_reserved
        or ip.is_unspecified
    )


def validate_url_for_ssrf(url: str):
    """Validate URL to reduce SSRF risk."""
    try:
        parsed = urlparse((url or "").strip())
        if parsed.scheme not in {"http", "https"}:
            return False, "Only http/https URLs are allowed"

        hostname = parsed.hostname
        if not hostname:
            return False, "URL hostname is required"

        # Direct IP path
        try:
            if _is_blocked_ip(hostname) and not _host_is_allowlisted(hostname):
                return False, "Private or reserved hosts are not allowed"
        except ValueError:
            # Hostname path: resolve and validate target addresses
            try:
                resolved_addresses = {
                    result[4][0] for result in socket.getaddrinfo(hostname, parsed.port or 80)
                }
            except socket.gaierror:
                return False, "Hostname cannot be resolved"

            for addr in resolved_addresses:
                if _is_blocked_ip(addr) and not _host_is_allowlisted(hostname):
                    return False, "URL resolves to a private or reserved host"

        return True, ""
    except Exception:
        return False, "Invalid URL format"


def _validate_chunks(chunks):
    if not isinstance(chunks, list):
        raise ValueError(f"Chunks must be a list, got {type(chunks).__name__}")
    for idx, item in enumerate(chunks):
        if not isinstance(item, str):
            raise ValueError(f"Chunk {idx} must be a string, got {type(item).__name__}")


def _write_chunks_json_gz(chunks, path):
    _validate_chunks(chunks)
    with gzip.open(path, "wt", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False)


def _load_chunks_safe(chunks_path, allow_legacy_pickle=False):
    """Load chunks safely and validate shape."""
    try:
        if chunks_path.endswith(".json.gz"):
            with gzip.open(chunks_path, "rt", encoding="utf-8") as f:
                chunks = json.load(f)
            _validate_chunks(chunks)
            return chunks

        if allow_legacy_pickle and (
            chunks_path.endswith(".pkl.gz") or chunks_path.endswith(".pkl")
        ):
            with (
                gzip.open(chunks_path, "rb")
                if chunks_path.endswith(".gz")
                else open(chunks_path, "rb")
            ) as f:
                chunks = pickle.load(f)
            _validate_chunks(chunks)
            return chunks

        raise ValueError("Unsupported chunk format")
    except Exception as exc:
        raise ValueError(f"Failed to load chunks: {exc}") from exc


def _migrate_legacy_chunks_to_json(kb_id):
    """Migrate legacy pickle chunk files to JSON gz format."""
    kb_folder = f"app/knowledge_bases/{kb_id}"
    target_path = os.path.join(kb_folder, CHUNKS_JSON_GZ)
    if os.path.exists(target_path):
        return target_path

    legacy_candidates = [
        os.path.join(kb_folder, LEGACY_CHUNKS_GZ),
        os.path.join(kb_folder, LEGACY_CHUNKS_RAW),
    ]
    for legacy_path in legacy_candidates:
        if not os.path.exists(legacy_path):
            continue
        chunks = _load_chunks_safe(legacy_path, allow_legacy_pickle=True)
        _write_chunks_json_gz(chunks, target_path)
        logger.warning("Migrated legacy chunks format for KB %s", kb_id)
        return target_path
    return None

def chunk_text(
    text_stream: Iterable[str],
    max_tokens: int = 800,
    overlap_tokens: int = 200,
    encoding_name: str = "cl100k_base"
) -> Iterator[str]:
    """
    Split iterable text segments into overlapping chunks sized by tokens.
    Falls back to character-based chunking when tiktoken isn't available.
    """
    if isinstance(text_stream, str):
        text_stream = [text_stream]
    if overlap_tokens >= max_tokens:
        raise ValueError("overlap_tokens must be smaller than max_tokens")

    if not tiktoken:
        print("[WARN] tiktoken not installed; falling back to character-based chunking")
        buffer = ""
        for segment in text_stream:
            if not segment:
                continue
            buffer += segment
            while len(buffer) >= max_tokens:
                chunk = buffer[:max_tokens]
                yield chunk
                buffer = buffer[max_tokens - overlap_tokens :]
        if buffer:
            yield buffer
        return

    encoding = tiktoken.get_encoding(encoding_name)
    token_buffer: List[int] = []
    for segment in text_stream:
        if not segment:
            continue
        segment_tokens = encoding.encode(segment)
        token_buffer.extend(segment_tokens)

        while len(token_buffer) >= max_tokens:
            chunk_tokens = token_buffer[:max_tokens]
            yield encoding.decode(chunk_tokens)
            token_buffer = token_buffer[max_tokens - overlap_tokens :]

    if token_buffer:
        yield encoding.decode(token_buffer)

def create_embeddings(chunks: List[str], batch_size: int = 32, provider: str = "openai", model: str = None) -> np.ndarray:
    """
    Create embeddings for text chunks in batches.
    
    Args:
        chunks: List of text chunks to embed
        batch_size: Number of chunks per API call
        provider: Embedding provider ("openai" or "gemini")
        model: Optional specific model name
        
    Returns:
        numpy array of embeddings
    """
    if not chunks:
        # Get default dimension for provider
        from app.config.provider_config import get_provider_metadata
        default_dim = get_provider_metadata(provider).get("embedding_dimensions", DEFAULT_EMBEDDING_DIM)
        return np.zeros((0, default_dim), dtype="float32")

    embeddings: List[np.ndarray] = []
    total_chunks = len(chunks)
    embedding_dim = None

    # For Gemini, we need to embed one at a time (no batch support in wrapper yet)
    # For OpenAI, we can batch but use individual calls for consistency
    for idx, chunk in enumerate(chunks, 1):
        try:
            vector = create_embedding(chunk, provider=provider, model=model)
            if embedding_dim is None:
                embedding_dim = vector.shape[0]
            embeddings.append(vector)
        except Exception as e:
            print(f"Error creating embedding for chunk: {e}")
            if embedding_dim is None:
                from app.config.provider_config import get_provider_metadata
                embedding_dim = get_provider_metadata(provider).get(
                    "embedding_dimensions", DEFAULT_EMBEDDING_DIM
                )
            embeddings.append(np.zeros(embedding_dim, dtype="float32"))

        print(f"🧠 Embedded {idx}/{total_chunks} chunks")

    if not embeddings:
        return np.zeros((0, embedding_dim or DEFAULT_EMBEDDING_DIM), dtype="float32")

    return np.vstack(embeddings)

def generate_ai_summary(text, title="", source_type="document"):
    """Generate an AI summary/description of the knowledge base content"""
    try:
        # Truncate text if too long (keep first 8000 characters for context)
        truncated_text = text[:8000] if len(text) > 8000 else text
        
        prompt = f"""Analyze the following {source_type} content and write a concise, helpful summary that describes what information is contained within. This summary will be used as a description to help AI agents understand what knowledge is available in this source.

Title: {title}
Content Type: {source_type}

Content:
{truncated_text}

Write a 2-3 sentence summary that describes:
1. What type of information this contains
2. Key topics or subjects covered
3. How this information might be useful for answering questions

Keep it concise but informative. Focus on the practical value and scope of the content."""

        response = _get_openai_client().chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that creates concise, informative summaries of knowledge base content."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,
            temperature=0.3
        )
        
        summary = response.choices[0].message.content.strip()
        return summary
        
    except Exception as e:
        print(f"Error generating AI summary: {e}")
        return f"Knowledge base containing {source_type} content related to {title}."

# Legacy web scraping functions removed - now using direct OpenAI approach for better results


def process_knowledge_base(kb_id, kb_type, source_path, generate_summary=False, embedding_provider="openai", embedding_model=None):
    """
    Process a knowledge base and create embeddings.
    
    Args:
        kb_id: Knowledge base ID
        kb_type: Type of source ("file", "url")
        source_path: Path to source file or URL
        generate_summary: Whether to generate AI summary
        embedding_provider: Provider for embeddings ("openai" or "gemini")
        embedding_model: Optional specific embedding model
    """
    try:
        # Check if this is an event knowledge base
        from app.config.knowledge_config import load_knowledge_config
        config = load_knowledge_config()
        is_event_kb = False
        kb_info = None
        event_source_text = None
        
        for kb in config.get('knowledge_bases', []):
            if kb.get('id') == kb_id:
                is_event_kb = kb.get('is_events', False)
                kb_info = kb
                break
        
        # Extract text based on type
        if kb_type == "file":
            if source_path.lower().endswith(".docx"):
                text_iterator = extract_text_from_docx(source_path)
            elif source_path.lower().endswith(".pdf"):
                text_iterator = extract_text_from_pdf(source_path)
            else:
                print(f"Unsupported file type: {source_path}")
                return False, None
        elif kb_type == "url":
            fetched_text = fetch_content_from_url(source_path)
            event_source_text = fetched_text
            text_iterator = [fetched_text] if fetched_text else []
        else:
            print(f"Unsupported knowledge base type: {kb_type}")
            return False, None

        if text_iterator is None:
            print(f"No text iterator available for {source_path}")
            return False, None
        
        segments_emitted = 0
        summary_char_limit = 8000
        summary_buffer: List[str] = []
        summary_chars = 0

        def streaming_segments() -> Iterator[str]:
            nonlocal segments_emitted, summary_chars
            for segment in text_iterator:
                if not segment:
                    continue
                segments_emitted += 1
                if generate_summary and summary_chars < summary_char_limit:
                    remaining = summary_char_limit - summary_chars
                    snippet = segment[:remaining]
                    summary_buffer.append(snippet)
                    summary_chars += len(snippet)
                yield segment

        chunk_iter = chunk_text(streaming_segments())
        chunks = list(chunk_iter)

        if segments_emitted == 0 or not chunks:
            print(f"No text chunks produced from {source_path}")
            return False, None
        
        print(f"🧩 Generated {len(chunks)} chunks for knowledge base {kb_id}")
        
        # Generate AI summary if requested
        ai_summary = None
        if generate_summary:
            title = kb_info.get('title', '') if kb_info else ""
            summary_text = "".join(summary_buffer)
            ai_summary = generate_ai_summary(summary_text, title, kb_type)
        
        # Create embeddings
        embeddings = create_embeddings(chunks, provider=embedding_provider, model=embedding_model)
        
        if embeddings.size == 0:
            print(f"No embeddings generated for {kb_id}")
            return False, ai_summary

        # Save checkpoints prior to index build
        kb_folder = f"app/knowledge_bases/{kb_id}"
        os.makedirs(kb_folder, exist_ok=True)
        
        index_path = f"{kb_folder}/index.faiss"
        chunks_path = f"{kb_folder}/{CHUNKS_JSON_GZ}"
        embeddings_path = f"{kb_folder}/embeddings.npy"
        
        _write_chunks_json_gz(chunks, chunks_path)
        np.save(embeddings_path, embeddings)

        # Create FAISS index
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        index.add(embeddings)
        
        faiss.write_index(index, index_path)
        
        print(f"Successfully processed knowledge base {kb_id}")
        return True, ai_summary
        
    except Exception as e:
        print(f"Error processing knowledge base {kb_id}: {e}")
        return False, None


def search_knowledge_base_with_embedding(kb_id, query_embedding, top_k=3):
    """
    Search a specific knowledge base with a pre-computed embedding.
    
    Args:
        kb_id: Knowledge base ID
        query_embedding: Pre-computed embedding vector (numpy array)
        top_k: Number of results to return
        
    Returns:
        List of tuples: (chunk_text, distance, kb_id)
        Distance is L2 distance (lower = more similar)
    """
    try:
        kb_folder = f"app/knowledge_bases/{kb_id}"
        index_path = f"{kb_folder}/index.faiss"
        chunks_path = f"{kb_folder}/{CHUNKS_JSON_GZ}"
        if not os.path.exists(chunks_path):
            migrated_path = _migrate_legacy_chunks_to_json(kb_id)
            if migrated_path:
                chunks_path = migrated_path
        
        if not os.path.exists(index_path) or not chunks_path or not os.path.exists(chunks_path):
            return []
        
        # Load index and chunks
        index = faiss.read_index(index_path)
        chunks = _load_chunks_safe(chunks_path, allow_legacy_pickle=False)
        
        # Ensure embedding is the right shape
        if query_embedding.ndim == 1:
            query_embedding = query_embedding.reshape(1, -1)
        
        # Search
        distances, indices = index.search(query_embedding, top_k)
        
        # Return results with distance and kb_id
        results = []
        for idx, dist_idx in enumerate(indices[0]):
            if dist_idx < len(chunks):
                chunk = chunks[dist_idx]
                distance = float(distances[0][idx])
                results.append((chunk, distance, kb_id))
        
        return results
        
    except Exception as e:
        print(f"Error searching knowledge base {kb_id}: {e}")
        return []

def search_knowledge_base(kb_id, query, top_k=3):
    """
    Search a specific knowledge base with a text query.
    Creates embedding using KB's provider (defaults to OpenAI for backward compatibility).
    
    Returns:
        List of tuples: (chunk_text, distance, kb_id)
        Distance is L2 distance (lower = more similar)
    """
    try:
        # Get KB info to determine provider
        from app.config.knowledge_config import load_knowledge_config
        config = load_knowledge_config()
        kb_info = None
        for kb in config.get('knowledge_bases', []):
            if kb.get('id') == kb_id:
                kb_info = kb
                break
        
        provider = kb_info.get('embedding_provider', 'openai') if kb_info else 'openai'
        model = kb_info.get('embedding_model') if kb_info else None
        
        # Create query embedding with appropriate provider
        query_embedding = create_embedding(query, provider=provider, model=model)
        
        # Use the embedding-based search
        return search_knowledge_base_with_embedding(kb_id, query_embedding, top_k)
        
    except Exception as e:
        print(f"Error searching knowledge base {kb_id}: {e}")
        return []